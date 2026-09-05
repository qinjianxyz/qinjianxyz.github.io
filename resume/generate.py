"""Build the public résumé PDF, editable DOCX, Markdown, and HTML from content.json.
Consumer: the personal website and job applications. Verify PDF text and one-page
layout after generation. Requires reportlab, python-docx, and pypdf.
"""
from pathlib import Path
import argparse
import html
import json
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
DATA = json.loads((ROOT / 'resume/content.json').read_text())
INK, BLUE, MUTED = '#1b262b', '#244cdf', '#4c585e'

def markdown(d):
    lines = [f"# {d['name']}", '', d['headline'], '',
             f"{d['location']} · {d['authorization']}",
             f"{d['phone']} · {d['email']}",
             f"https://{d['linkedin']} · https://{d['website']}", '', '## Experience']
    for job in d['experience']:
        lines += ['', f"### {job['company']}", f"{job['role']} | {job['dates']}", '']
        lines += ['- ' + text for text in job['bullets']]
    lines += ['', '## Selected Projects']
    for p in d['projects']:
        lines += ['', f"### {p['name']}", p['context'], '', p['text']]
    e = d['education']
    lines += ['', '## Education', '', f"### {e['school']}", f"{e['degree']} | {e['dates']}", e['honors'], '', '## Technical Skills', '']
    lines += [f'**{k}:** {v}  ' for k,v in d['skills']]
    return '\n'.join(lines) + '\n'

def pdf(d, font_dir):
    for key, file in [('Resume','Arial.ttf'), ('Resume-Bold','Arial Bold.ttf')]:
        pdfmetrics.registerFont(TTFont(key,str(font_dir/file)))
    pdfmetrics.registerFontFamily('Resume',normal='Resume',bold='Resume-Bold',italic='Resume',boldItalic='Resume-Bold')
    base = dict(fontName='Resume',fontSize=10,leading=13,textColor=HexColor(INK))
    styles = {
      'body':ParagraphStyle('body',**base,spaceAfter=5),
      'bullet':ParagraphStyle('bullet',**base,leftIndent=10,firstLineIndent=-8,spaceAfter=4),
      'name':ParagraphStyle('name',fontName='Resume-Bold',fontSize=27,leading=30,textColor=HexColor(INK),spaceAfter=4),
      'headline':ParagraphStyle('headline',fontName='Resume-Bold',fontSize=10.5,leading=14,textColor=HexColor(BLUE),spaceAfter=5),
      'meta':ParagraphStyle('meta',fontName='Resume',fontSize=8.5,leading=11.5,textColor=HexColor(MUTED),spaceAfter=3),
      'section':ParagraphStyle('section',fontName='Resume-Bold',fontSize=9.2,leading=12,textColor=HexColor(BLUE),spaceBefore=8,spaceAfter=5,keepWithNext=True),
      'company':ParagraphStyle('company',fontName='Resume-Bold',fontSize=10.5,leading=13,textColor=HexColor(INK),spaceAfter=2,keepWithNext=True),
      'role':ParagraphStyle('role',fontName='Resume',fontSize=8.7,leading=11.5,textColor=HexColor(MUTED),spaceAfter=5,keepWithNext=True),
      'skill':ParagraphStyle('skill',fontName='Resume',fontSize=8.7,leading=11.5,textColor=HexColor(INK),spaceAfter=2),
    }
    story=[]
    def add(text,kind='body',raw=False):
        story.append(Paragraph(text if raw else html.escape(text),styles[kind]))
    add(d['name'],'name');add(d['headline'],'headline')
    add(f"{d['location']} · {d['authorization']}",'meta')
    add(f"{d['phone']} · <link href='mailto:{d['email']}'>{d['email']}</link>",'meta',True)
    add(f"<link href='https://{d['linkedin']}'>{d['linkedin']}</link> · <link href='https://{d['website']}'>{d['website']}</link>",'meta',True)
    story.append(Spacer(1,3))
    def section(text):
        add(text.upper(),'section')
        story.append(HRFlowable(width='100%',thickness=.5,color=HexColor('#d9dfdc'),spaceAfter=5))
    section('Experience')
    for job in d['experience']:
        add(job['company'],'company');add(f"{job['role']} | {job['dates']}",'role')
        for b in job['bullets']:add('• '+b,'bullet')
        story.append(Spacer(1,3))
    section('Selected Projects')
    for p in d['projects']:
        add(f"{p['name']} <font size='8.5' color='{MUTED}'> / {html.escape(p['context'])}</font>",'company',True)
        add(p['text'])
    section('Education');e=d['education']
    add(e['school'],'company');add(f"{e['degree']} | {e['dates']}",'meta');add(e['honors'],'meta')
    section('Technical Skills')
    for k,v in d['skills']:add(f'<b>{html.escape(k)}:</b> {html.escape(v)}','skill',True)
    out=ROOT/'media/.Ray_Qin_Resume.build.pdf'
    doc=SimpleDocTemplate(str(out),pagesize=(612,792),leftMargin=42,rightMargin=42,topMargin=34,bottomMargin=32,title='Ray Qin — Senior Software Engineer',author='Ray Qin')
    doc.build(story)
    reader=PdfReader(out)
    assert len(reader.pages)==1, f'Resume overflowed: {len(reader.pages)} pages'
    text=reader.pages[0].extract_text()
    for expected in [d['name'],'25M+','10K+','Eightfold','Ground Truth','3.9/4.0','OpenTelemetry']:
        assert expected in text, f'Missing extracted content: {expected}'
    assert len(reader.pages[0].get('/Annots',[]))>=3,'Missing contact hyperlinks'
    out.replace(ROOT/'media/Ray_Qin_Resume.pdf')
    (ROOT/'media/Ray_Resume.pdf').write_bytes((ROOT/'media/Ray_Qin_Resume.pdf').read_bytes())
    return len(text.split())

def docx(d):
    doc=Document();sec=doc.sections[0]
    sec.page_width=Inches(8.5);sec.page_height=Inches(11)
    sec.top_margin=Inches(.47);sec.bottom_margin=Inches(.44)
    sec.left_margin=sec.right_margin=Inches(.58)
    normal=doc.styles['Normal'];normal.font.name='Arial';normal.font.size=Pt(10)
    normal.font.color.rgb=RGBColor.from_string('1B262B')
    normal.paragraph_format.space_after=Pt(4);normal.paragraph_format.line_spacing=1.05
    for name,size,col in [('Title',27,'1B262B'),('Heading 1',9.2,'244CDF'),('Heading 2',10.5,'1B262B')]:
        st=doc.styles[name];st.font.name='Arial';st.font.size=Pt(size);st.font.bold=True;st.font.color.rgb=RGBColor.from_string(col)
        st.paragraph_format.space_before=Pt(8 if name=='Heading 1' else 2);st.paragraph_format.space_after=Pt(4)
    doc.add_paragraph(d['name'],'Title')
    p=doc.add_paragraph();r=p.add_run(d['headline']);r.bold=True;r.font.color.rgb=RGBColor.from_string('244CDF')
    for text in [f"{d['location']} · {d['authorization']}",f"{d['phone']} · {d['email']}",f"{d['linkedin']} · {d['website']}"]:
        p=doc.add_paragraph(text);p.paragraph_format.space_after=Pt(2)
        for r in p.runs:r.font.size=Pt(8.5)
    doc.add_paragraph('EXPERIENCE','Heading 1')
    for j in d['experience']:
        doc.add_paragraph(j['company'],'Heading 2');p=doc.add_paragraph(f"{j['role']} | {j['dates']}");p.paragraph_format.keep_with_next=True
        for r in p.runs:r.font.size=Pt(8.7)
        for b in j['bullets']:
            p=doc.add_paragraph(b,'List Bullet');p.paragraph_format.left_indent=Inches(.14);p.paragraph_format.first_line_indent=Inches(-.12)
    doc.add_paragraph('SELECTED PROJECTS','Heading 1')
    for p in d['projects']:
        doc.add_paragraph(p['name']+' / '+p['context'],'Heading 2');doc.add_paragraph(p['text'])
    doc.add_paragraph('EDUCATION','Heading 1');e=d['education']
    doc.add_paragraph(e['school'],'Heading 2');doc.add_paragraph(e['degree']+' | '+e['dates']);doc.add_paragraph(e['honors'])
    doc.add_paragraph('TECHNICAL SKILLS','Heading 1')
    for k,v in d['skills']:
        p=doc.add_paragraph();p.add_run(k+': ').bold=True;p.add_run(v)
        p.paragraph_format.space_after=Pt(2)
        for r in p.runs:r.font.size=Pt(8.7)
    doc.core_properties.author='Ray Qin';doc.core_properties.title='Ray Qin — Senior Software Engineer'
    doc.save(ROOT/'media/Ray_Qin_Resume.docx')

def webpage(d):
    esc = html.escape
    index = (ROOT/'index.html').read_text()
    header = index[index.index('  <header class="site-nav">'):index.index('  </header>')+len('  </header>')]
    header = header.replace('href="resume.html"', 'href="resume.html" aria-current="page"')
    footer = index[index.index('  <footer class="footer">'):index.index('  <script src="assets/site.js"')]
    fonts=index[index.index('  <link rel="preconnect"'):index.index('  <link rel="stylesheet"')]
    sections=[]
    for job in d['experience']:
        bullets=''.join('<li>'+esc(b)+'</li>' for b in job['bullets'])
        sections.append(f"<article><h3>{esc(job['company'])}</h3><p class='resume-meta'>{esc(job['role'])} · {esc(job['dates'])}</p><ul>{bullets}</ul></article>")
    projects=''.join(f"<article><h3>{esc(p['name'])}</h3><p class='resume-meta'>{esc(p['context'])}</p><p>{esc(p['text'])}</p></article>" for p in d['projects'])
    skills=''.join(f'<div><dt>{esc(k)}</dt><dd>{esc(v)}</dd></div>' for k,v in d['skills'])
    e=d['education']
    page=f'''<!doctype html>
<html lang="en"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>Ray Qin Résumé | Senior Software Engineer &amp; Applied AI</title>
<meta name="description" content="Ray Qin's engineering experience: production platforms at ZipRecruiter, applied AI, retrieval systems, and independent projects. Read online or download PDF and Word.">
<meta name="author" content="Ray Qin"><link rel="canonical" href="https://qinjianxyz.github.io/resume.html"><link rel="icon" href="favicon.ico">
<meta property="og:title" content="Ray Qin | Engineering résumé"><meta property="og:type" content="website"><meta property="og:url" content="https://qinjianxyz.github.io/resume.html"><meta property="og:image" content="https://qinjianxyz.github.io/media/ray-qin-og-v2.png">
<meta name="twitter:card" content="summary_large_image"><meta name="theme-color" content="#fafbf8">
{fonts}<link rel="stylesheet" href="assets/site.css"><link rel="stylesheet" href="assets/resume.css">
</head><body><a class="skip-link" href="#main">Skip to content</a>
{header}
<main id="main" class="resume-main"><div class="wrap">
<div class="resume-toolbar"><p class="kicker">Experience / Résumé</p><div><a class="btn btn-primary" href="media/Ray_Qin_Resume.pdf" download>Download PDF <span aria-hidden="true">↓</span></a><a class="btn btn-secondary" href="media/Ray_Qin_Resume.docx" download>Word version <span aria-hidden="true">↓</span></a></div></div>
<div class="resume-sheet"><header class="resume-heading"><h1>{esc(d['name'])}</h1><p class="resume-title">{esc(d['headline'])}</p><p>{esc(d['location'])} · {esc(d['authorization'])}</p><div class="resume-contact"><a href="mailto:{d['email']}">{d['email']}</a><a href="tel:+14156046032">{d['phone']}</a><a href="https://{d['linkedin']}">LinkedIn ↗</a><a href="https://{d['website']}">Portfolio ↗</a></div></header>
<section aria-labelledby="resume-experience"><h2 id="resume-experience">Experience</h2>{''.join(sections)}</section>
<section aria-labelledby="resume-projects"><h2 id="resume-projects">Selected projects</h2>{projects}</section>
<section aria-labelledby="resume-education"><h2 id="resume-education">Education</h2><article><h3>{esc(e['school'])}</h3><p>{esc(e['degree'])}</p><p class="resume-meta">{esc(e['dates'])} · {esc(e['honors'])}</p></article></section>
<section aria-labelledby="resume-skills"><h2 id="resume-skills">Technical skills</h2><dl class="resume-skills">{skills}</dl></section>
</div><p class="resume-endnote">Prefer to see the work? <a href="index.html#proof">Explore selected projects ↗</a></p>
</div></main>{footer}<script src="assets/site.js" defer></script></body></html>'''
    (ROOT/'resume.html').write_text(page)

if __name__=='__main__':
    parser=argparse.ArgumentParser();parser.add_argument('--font-dir',type=Path,default=Path('/System/Library/Fonts/Supplemental'));args=parser.parse_args()
    (ROOT/'resume/Ray_Qin_Resume.md').write_text(markdown(DATA))
    count=pdf(DATA,args.font_dir);docx(DATA);webpage(DATA)
    print(f'PASS: one-page PDF; {count} extracted words; contact hyperlinks; DOCX, Markdown, and web page generated.')
